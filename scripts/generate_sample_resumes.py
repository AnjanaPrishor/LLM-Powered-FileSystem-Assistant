"""
scripts/generate_sample_resumes.py
----------------------------------
Idempotent generator for the .pdf and .docx dummy resumes in sample_data/.

Plain-text resumes (.txt, .md) are committed directly; this script only
creates the binary formats so we don't have to check them into git.

Run:
    py scripts/generate_sample_resumes.py
"""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
SAMPLE_DIR = ROOT / "sample_data"
SAMPLE_DIR.mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# Resume content (kept small so the file stays readable)
# ---------------------------------------------------------------------------

MICHAEL_CHEN_PDF = [
    "Michael Chen",
    "Senior Java Backend Developer",
    "",
    "Location: Seattle, WA",
    "Email: michael.chen@example.com",
    "",
    "Experience:",
    "- 9 years building JVM microservices at scale",
    "- Led migration from Spring MVC to Spring Boot 3 across 15 services",
    "- Optimized Kafka consumer throughput by 4x at PayNimbus (2024)",
    "",
    "Skills: Java, Kotlin, Spring Boot, Kafka, PostgreSQL, Docker, AWS,",
    "gRPC, JUnit, some Python for tooling scripts.",
    "",
    "Education:",
    "- M.S. Computer Science, University of Washington, 2016",
]

LISA_WONG_PDF = [
    "Lisa Wong",
    "Cybersecurity Analyst",
    "",
    "Location: Singapore",
    "Email: lisa.wong@example.com",
    "",
    "Experience:",
    "- 6 years in SOC and incident response at FinShield Bank",
    "- Built SIEM detection rules in Splunk and Elastic",
    "- Ran phishing simulations for a 4,000-person org",
    "",
    "Skills: Splunk, Elastic SIEM, Python, Wireshark, MITRE ATT&CK,",
    "AWS GuardDuty, Burp Suite, threat modeling.",
    "",
    "Certifications: CISSP, OSCP, AWS Security Specialty",
    "",
    "Education:",
    "- B.Eng. Information Security, NUS, 2019",
]

SARA_KIM_DOCX = [
    ("heading", "Sara Kim", 1),
    ("para", "Product Designer (UX/UI)"),
    ("para", "Location: Seoul, South Korea"),
    ("para", "Email: sara.kim@example.com"),
    ("heading", "Summary", 2),
    ("para",
     "Product designer with 5 years shipping consumer mobile apps. "
     "Comfortable pairing with engineers and running usability tests."),
    ("heading", "Experience", 2),
    ("para",
     "- Senior Product Designer, KakaoWorks (2023 - present): "
     "led redesign of onboarding flow, cutting drop-off by 22%."),
    ("para",
     "- Product Designer, LineArt Studio (2020 - 2023): "
     "shipped 3 iOS apps, ran 40+ user interviews."),
    ("heading", "Skills", 2),
    ("para", "Figma, Sketch, prototyping, design systems, HTML/CSS, "
            "basic JavaScript, some Python for automating Figma exports."),
    ("heading", "Education", 2),
    ("para", "B.F.A. Visual Design, Hongik University, 2020"),
]

OMAR_HASSAN_DOCX = [
    ("heading", "Omar Hassan", 1),
    ("para", "Cloud Solutions Architect"),
    ("para", "Location: Dubai, UAE"),
    ("para", "Email: omar.hassan@example.com"),
    ("heading", "Summary", 2),
    ("para",
     "Cloud architect with 11 years designing multi-region AWS and Azure "
     "workloads for banking and telecom clients."),
    ("heading", "Experience", 2),
    ("para",
     "- Principal Cloud Architect, GulfCloud Consulting (2022 - present): "
     "designed an active-active AWS deployment across 3 regions for a "
     "national bank."),
    ("para",
     "- Senior Cloud Engineer, Etisalat Digital (2018 - 2022): "
     "led the Kubernetes platform team, on-call rotation for 200+ services."),
    ("heading", "Skills", 2),
    ("para", "AWS, Azure, Kubernetes, Terraform, Python, Go, networking, "
            "cost optimization, FinOps, Well-Architected reviews."),
    ("heading", "Certifications", 2),
    ("para", "AWS Solutions Architect Professional, Azure Solutions "
            "Architect Expert, CKA."),
    ("heading", "Education", 2),
    ("para", "B.Sc. Computer Engineering, American University of Sharjah, 2014"),
]


# ---------------------------------------------------------------------------
# Writers
# ---------------------------------------------------------------------------

def write_pdf(path: Path, lines: list[str]) -> None:
    """Render a small PDF using reportlab (required for these files)."""
    try:
        from reportlab.pdfgen import canvas
    except ImportError as exc:  # pragma: no cover
        raise SystemExit(
            "reportlab is required to generate the sample PDFs. "
            "Install it with:  py -m pip install reportlab"
        ) from exc

    c = canvas.Canvas(str(path))
    text = c.beginText(72, 750)
    text.setFont("Helvetica", 11)
    for line in lines:
        text.textLine(line)
    c.drawText(text)
    c.showPage()
    c.save()


def write_docx(path: Path, blocks: list[tuple]) -> None:
    """Render a small DOCX using python-docx."""
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise SystemExit(
            "python-docx is required to generate the sample DOCX files. "
            "Install it with:  py -m pip install python-docx"
        ) from exc

    doc = docx.Document()
    for block in blocks:
        kind = block[0]
        if kind == "heading":
            _, text, level = block
            doc.add_heading(text, level=level)
        elif kind == "para":
            _, text = block
            doc.add_paragraph(text)
        else:  # pragma: no cover
            raise ValueError(f"Unknown block type: {kind}")
    doc.save(str(path))


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    generated: list[Path] = []

    pdf_targets = [
        (SAMPLE_DIR / "resume_michael_chen.pdf", MICHAEL_CHEN_PDF),
        (SAMPLE_DIR / "resume_lisa_wong.pdf",    LISA_WONG_PDF),
    ]
    docx_targets = [
        (SAMPLE_DIR / "resume_sara_kim.docx",    SARA_KIM_DOCX),
        (SAMPLE_DIR / "resume_omar_hassan.docx", OMAR_HASSAN_DOCX),
    ]

    for path, lines in pdf_targets:
        write_pdf(path, lines)
        generated.append(path)

    for path, blocks in docx_targets:
        write_docx(path, blocks)
        generated.append(path)

    print("Generated sample resumes:")
    for p in generated:
        print(f"  {p.relative_to(ROOT)} ({p.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
