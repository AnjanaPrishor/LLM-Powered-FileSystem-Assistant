"""
tests/test_fs_tools.py

Small, no-framework smoke test for fs_tools. Generates a real .pdf and
.docx on the fly (so we don't need any binary test assets in the repo),
then exercises every tool and prints the results.

Run:  py tests/test_fs_tools.py
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

# Make sure we can import fs_tools when run from anywhere
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from fs_tools import read_file, list_files, write_file, search_in_file  # noqa: E402

SAMPLE_DIR = ROOT / "sample_data"
SAMPLE_DIR.mkdir(exist_ok=True)


def make_sample_pdf() -> Path:
    """Create sample_data/resume.pdf using pypdf's writer."""
    from pypdf import PdfWriter
    from pypdf.generic import RectangleObject

    path = SAMPLE_DIR / "resume.pdf"

    # pypdf can't easily draw text from scratch; use reportlab if present,
    # otherwise fall back to a tiny hand-rolled PDF with embedded text.
    try:
        from reportlab.pdfgen import canvas  # type: ignore
        c = canvas.Canvas(str(path))
        text = c.beginText(72, 750)
        for line in [
            "John Smith",
            "Machine Learning Engineer",
            "",
            "Experience:",
            "- 4 years building LLM-powered applications",
            "- Strong Python, PyTorch, and FastAPI skills",
            "",
            "Skills: Python, LLMs, PyTorch, Docker, AWS",
        ]:
            text.textLine(line)
        c.drawText(text)
        c.showPage()
        c.save()
    except ImportError:
        # Minimal fallback: write a *very* small valid PDF containing the
        # word "Python" so search_in_file has something to find.
        path.write_bytes(
            b"%PDF-1.4\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n"
            b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 300 144]"
            b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
            b"4 0 obj<</Length 55>>stream\n"
            b"BT /F1 18 Tf 20 100 Td (Python LLM Resume) Tj ET\n"
            b"endstream endobj\n"
            b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
            b"xref\n0 6\n0000000000 65535 f \n"
            b"0000000009 00000 n \n0000000053 00000 n \n"
            b"0000000098 00000 n \n0000000181 00000 n \n"
            b"0000000262 00000 n \n"
            b"trailer<</Size 6/Root 1 0 R>>\n"
            b"startxref\n320\n%%EOF"
        )
    return path


def make_sample_docx() -> Path:
    """Create sample_data/resume.docx using python-docx."""
    import docx
    path = SAMPLE_DIR / "resume.docx"
    doc = docx.Document()
    doc.add_heading("Alice Nguyen", level=1)
    doc.add_paragraph("Full-Stack Engineer")
    doc.add_paragraph("Experience:")
    doc.add_paragraph("- 5 years building web apps with React and Node")
    doc.add_paragraph("- Some Python for data pipelines")
    doc.add_paragraph("Skills: JavaScript, TypeScript, Python, SQL")
    doc.save(str(path))
    return path


def banner(title: str) -> None:
    line = "=" * 60
    print(f"\n{line}\n{title}\n{line}")


def dump(obj) -> None:
    print(json.dumps(obj, indent=2, default=str))


# ---------------------------------------------------------------------------
# Tiny test-runner: tracks pass/fail counts and prints a summary.
# ---------------------------------------------------------------------------

class TestRunner:
    def __init__(self) -> None:
        self.passed: list[str] = []
        self.failed: list[tuple[str, str]] = []  # (name, reason)

    def run(self, name: str, action, check, display=None) -> None:
        """
        name    : test label (e.g. "TEST 1: list_files(...)")
        action  : zero-arg callable that returns the tool result
        check   : callable(result) -> (ok: bool, reason: str)
        display : optional callable(result) -> printable object
                  (defaults to printing the raw result)
        """
        banner(name)
        try:
            result = action()
        except Exception as exc:
            print(f"[EXCEPTION] {exc}")
            self.failed.append((name, f"exception: {exc}"))
            return

        dump(display(result) if display else result)

        try:
            ok, reason = check(result)
        except Exception as exc:
            ok, reason = False, f"check raised: {exc}"

        if ok:
            print("[PASS]")
            self.passed.append(name)
        else:
            print(f"[FAIL] {reason}")
            self.failed.append((name, reason))

    def summary(self) -> int:
        total = len(self.passed) + len(self.failed)
        banner("SUMMARY")
        print(f"Total : {total}")
        print(f"Passed: {len(self.passed)}")
        print(f"Failed: {len(self.failed)}")
        if self.failed:
            print("\nFailed tests:")
            for name, reason in self.failed:
                print(f"  - {name}\n      reason: {reason}")
        return 0 if not self.failed else 1


# ---------------------------------------------------------------------------
# Per-test expectations (what "success" means for each case)
# ---------------------------------------------------------------------------

def expect_success_dict(result):
    if isinstance(result, dict) and result.get("success") is True:
        return True, ""
    return False, f"expected success=True dict, got: {result!r}"


def expect_failure_dict(result):
    if isinstance(result, dict) and result.get("success") is False and result.get("error"):
        return True, ""
    return False, f"expected success=False dict with error, got: {result!r}"


def expect_nonempty_list(result):
    if not isinstance(result, list):
        return False, f"expected list, got {type(result).__name__}"
    if not result:
        return False, "expected non-empty list"
    # First element must not be an error envelope
    if isinstance(result[0], dict) and result[0].get("success") is False:
        return False, f"list contains error envelope: {result[0]!r}"
    return True, ""


def expect_search_hits(min_hits: int):
    def _check(result):
        ok, reason = expect_success_dict(result)
        if not ok:
            return ok, reason
        count = result.get("match_count", 0)
        if count < min_hits:
            return False, f"expected >= {min_hits} matches, got {count}"
        return True, ""
    return _check


def truncate_content(result):
    if isinstance(result, dict) and "content" in result:
        return {**result, "content": (result.get("content") or "")[:120] + "..."}
    return result


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    banner("SETUP: creating sample PDF and DOCX")
    pdf_path = make_sample_pdf()
    docx_path = make_sample_docx()
    print("created:", pdf_path.name, "and", docx_path.name)

    runner = TestRunner()

    runner.run(
        "TEST 1: list_files(sample_data)",
        lambda: list_files(str(SAMPLE_DIR)),
        expect_nonempty_list,
    )

    runner.run(
        "TEST 2: list_files(sample_data, extension='pdf')",
        lambda: list_files(str(SAMPLE_DIR), extension="pdf"),
        lambda r: (
            expect_nonempty_list(r)[0] and all(f.get("extension") == ".pdf" for f in r),
            "expected only .pdf files in filtered result",
        ),
    )

    runner.run(
        "TEST 3: read_file(resume.txt)",
        lambda: read_file(str(SAMPLE_DIR / "resume.txt")),
        expect_success_dict,
        display=truncate_content,
    )

    runner.run(
        "TEST 4: read_file(resume.pdf)",
        lambda: read_file(str(pdf_path)),
        expect_success_dict,
        display=truncate_content,
    )

    runner.run(
        "TEST 5: read_file(resume.docx)",
        lambda: read_file(str(docx_path)),
        expect_success_dict,
        display=truncate_content,
    )

    runner.run(
        "TEST 6: write_file(output/notes.txt)",
        lambda: write_file(str(ROOT / "output" / "notes.txt"),
                           "Hello from write_file!\nSecond line."),
        expect_success_dict,
    )

    runner.run(
        "TEST 7: search_in_file(resume.txt, 'python') [case-insensitive]",
        lambda: search_in_file(str(SAMPLE_DIR / "resume.txt"), "python"),
        expect_search_hits(1),
    )

    runner.run(
        "TEST 8: search_in_file(resume.docx, 'python')",
        lambda: search_in_file(str(docx_path), "python"),
        expect_search_hits(1),
    )

    runner.run(
        "TEST 9: ERROR CASE - read_file(does_not_exist.txt)",
        lambda: read_file(str(SAMPLE_DIR / "does_not_exist.txt")),
        expect_failure_dict,
    )

    unsupported = SAMPLE_DIR / "unsupported.xyz"
    unsupported.write_text("nope")
    runner.run(
        "TEST 10: ERROR CASE - read_file(unsupported.xyz)",
        lambda: read_file(str(unsupported)),
        expect_failure_dict,
    )

    return runner.summary()


if __name__ == "__main__":
    sys.exit(main())
